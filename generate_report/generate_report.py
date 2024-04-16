import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Cm
from docx.oxml.shared import OxmlElement, qn
from sql.support_degree_database import query_ipv6_records
from sql.ConnectDB import get_mysql_conn


def set_indent(paragraph, left_indent=None, right_indent=None, first_line_indent=None, hanging_indent=None):
    """设置缩进，单位为字符

    :param paragraph: 某段落
    :param left_indent: 左缩进
    :param right_indent: 右缩进
    :param first_line_indent: 首行缩进
    :param hanging_indent: 悬挂缩进
    """
    assert not all([first_line_indent, hanging_indent]), '首行缩进与悬挂缩进不可同时设置'
    pPr = paragraph._element.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    if left_indent:
        ind.set(qn('w:leftChars'), str(left_indent * 100))
    if right_indent:
        ind.set(qn('w:rightChars'), str(right_indent * 100))
    if first_line_indent:
        ind.set(qn('w:firstLineChars'), str(first_line_indent * 100))
    if hanging_indent:
        ind.set(qn('w:hangingChars'), str(hanging_indent * 100))
    pPr.append(ind)


def set_table_attr(table):
    # for row in table.rows:  # 统一行高
    #     row.height = Cm(1.0)
    #
    # for column in table.columns:  # 统一列宽
    #     # column.width = Cm(3.0)  # 不起作用
    #     for cell in column.cells:
    #         cell.width = Cm(3.0)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 表居中

    for cell in table.row_cells(0):  # 第一行单元格
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 垂直对齐，居中

    for column in table.columns:  # 所有列
        for cell in column.cells:  # 所有单元格
            for paragraph in cell.paragraphs:  # 所有段落
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 水平对齐，居中


############################################## 1.获取需要填充的参数
domain = 'www.neu.edu.cn'
start_time = datetime.date.today()
resolved = True


def generate_report(domain):
    # 从数据库中获取结果，并且进行格式转换
    result = query_ipv6_records(conn=get_mysql_conn(), domain='www.google.com')
    print(result)
    domain = result[1]
    resolved = result[2]
    accessed = result[3]
    support_degree = str(result[4])
    connectivity = result[5]
    resolve_delay = str(result[8])
    TCP_establishment_resolution_delay = str(result[9])
    server_responds_first_packet_delay = str(result[10])
    server_responds_first_page_delay = str(result[11])
    access_stability = result[12]
    ipv6_authorization_system = result[13]
    start_time = str(result[14])
    # print(start_time)
    end_time = str(result[15])

    document = Document()
    ############################################## 2.设置报告标题
    document.add_heading('网站 IPv6 支持度检测报告', level=0)

    paragraph = document.add_paragraph('')
    run = paragraph.add_run('网站域名: ')
    font = run.font
    font.name = '宋体'
    font.bold = True
    run = paragraph.add_run(domain)
    font = run.font
    font.name = 'Times New Roman'
    font.bold = True
    font.underline = True

    paragraph = document.add_paragraph('')
    run = paragraph.add_run('支持IPV6: ')
    if accessed:
        support = '   支持    '
    else:
        support = '   不支持    '
    font = run.font
    font.name = '宋体'
    font.bold = True
    run = paragraph.add_run(support)
    font = run.font
    font.name = 'Times New Roman'
    font.bold = True
    font.underline = True

    paragraph = document.add_paragraph('')
    run = paragraph.add_run('检测时间: ')
    font = run.font
    font.name = '宋体'
    font.bold = True
    run = paragraph.add_run(start_time)
    font = run.font
    font.name = 'Times New Roman'
    font.bold = True
    font.underline = True

    ############################################## 3.标题1及下面内容
    document.add_heading('1.检测依据', level=1)
    content = ('网站IPv6支持度评测指标与测试方法是基于中华人民共和国通信行业标准YD/T3118-2016《网站IPv6支持度评测指标与测试方法》。该标准规定了网站IPv6支持度评测指标与测试方法，适用于IPv6'
               '网络环境下的Web网站，以及IPv6和IPv4共存网络环境下支持双栈模式的Web网站。')
    paragraph = document.add_paragraph()
    set_indent(paragraph, left_indent=0, right_indent=0, first_line_indent=2)
    run = paragraph.add_run(content)
    font = run.font
    font.name = '宋体'

    ############################################## 4.标题2及下面内容
    # 根据ipv6的支持度数据填充到一个表格中
    document.add_heading('2.检测报告', level=1)
    table = document.add_table(13, 2, style='Light Shading Accent 1')
    heading_cells = table.rows[0].cells
    heading_cells[0].text = '域名'
    heading_cells[1].text = domain
    cells = table.rows[1].cells
    cells[0].text = '是否可以被解析'
    if resolved:
        cells[1].text = '可以被解析'
    else:
        cells[1].text = '不可以被解析'
    cells = table.rows[2].cells
    cells[0].text = '是否可以被访问'
    if accessed:
        cells[1].text = '可以被访问'
    else:
        cells[1].text = '不可以被访问'
    cells = table.rows[3].cells
    cells[0].text = '支持度'
    cells[1].text = support_degree
    cells = table.rows[4].cells
    cells[0].text = '连通性'
    cells[1].text = connectivity
    cells = table.rows[5].cells
    cells[0].text = '域名解析时延'
    cells[1].text = resolve_delay + 's'
    cells = table.rows[6].cells
    cells[0].text = 'TCP建立解析时延指标'
    cells[1].text = TCP_establishment_resolution_delay + 's'
    cells = table.rows[7].cells
    cells[0].text = '服务器相应首包时延指标'
    cells[1].text = server_responds_first_packet_delay + 's'
    cells = table.rows[8].cells
    cells[0].text = '服务器相应首页时延指标'
    cells[1].text = server_responds_first_page_delay + 's'
    cells = table.rows[9].cells
    cells[0].text = 'ipv6访问稳定性'
    cells[1].text = access_stability
    cells = table.rows[10].cells
    cells[0].text = '是否具有ipv6授权体系'
    if ipv6_authorization_system:
        cells[1].text = '具有'
    else:
        cells[1].text = '不具有'

    cells = table.rows[11].cells
    cells[0].text = '计算开始时间'
    cells[1].text = start_time
    cells = table.rows[12].cells
    cells[0].text = '计算结束时间'
    cells[1].text = end_time
    set_table_attr(table)

    document.add_heading('3.不支持IPV6的链接', level=1)
    document.add_heading('3.1 不支持IPV6的二级链接', level=2)
    secondary_links = []
    table = document.add_table(1, 2, style='Light Shading Accent 1')
    heading_cells = table.rows[0].cells
    heading_cells[0].text = '二级链接'
    heading_cells[1].text = '是否支持ipv6'

    for link in secondary_links:
        cells = table.add_row().cells
        cells[0].text = link[0]
        cells[1].text = '不支持'

    document.add_heading('3.2 不支持IPV6的三级链接', level=2)
    tertiary_links = []
    table = document.add_table(1, 2, style='Light Shading Accent 1')
    heading_cells = table.rows[0].cells
    heading_cells[0].text = '三级链接'
    heading_cells[1].text = '是否支持ipv6'

    for link in secondary_links:
        cells = table.add_row().cells
        cells[0].text = link[0]
        cells[1].text = '不支持'

    document.save('generate_report.docx')


if __name__ == '__main__':
    generate_report(domain='www.google.com')
